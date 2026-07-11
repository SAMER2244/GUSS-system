"""
report_generator.py  —  V3.2 Smart Auditing (Logging + Exceptions)
=================================================================
ينشئ ملف Word (.docx) احترافيًا باللغة العربية واتجاه RTL.
يتضمن قسم التدقيق المقارن (الخطة المعتمدة vs الواقع التنفيذي).
"""

from __future__ import annotations

import re
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

import subprocess
import platform
import shutil
import config as cfg
from data_parser import OfficeData, get_task_statistics
from logger import get_logger
from exceptions import ReportGenerationError

_log = get_logger("report")


# ─── RTL / Bidi Helpers ──────────────────────────────────────────────────────
def _set_rtl(paragraph) -> None:
    """يضبط اتجاه الفقرة RTL ولغة Arabic."""
    pPr = paragraph._p.get_or_add_pPr()

    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    pPr.append(bidi)

    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "right")
    pPr.append(jc)


def _set_rtl_run(run) -> None:
    """يضبط خصائص RTL على مستوى الـ run."""
    rPr = run._r.get_or_add_rPr()
    rtl = OxmlElement("w:rtl")
    rtl.set(qn("w:val"), "1")
    rPr.append(rtl)

    lang = OxmlElement("w:lang")
    lang.set(qn("w:val"), "ar-SA")
    lang.set(qn("w:bidi"), "ar-SA")
    rPr.append(lang)


def _make_rtl_paragraph(doc: Document, text: str, style: str = "Normal",
                         bold: bool = False, size_pt: int = 12,
                         color: RGBColor | None = None,
                         align: str = "right") -> Any:
    """ينشئ فقرة RTL صحيحة ومُنسَّقة."""
    p = doc.add_paragraph(style=style)
    _set_rtl(p)

    # محاذاة
    align_map = {
        "right":   WD_ALIGN_PARAGRAPH.RIGHT,
        "center":  WD_ALIGN_PARAGRAPH.CENTER,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    p.alignment = align_map.get(align, WD_ALIGN_PARAGRAPH.RIGHT)

    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size_pt)
    if color:
        run.font.color.rgb = color
    # خط عربي
    run.font.name = cfg.REPORT_FONT
    run._r.rPr.rFonts.set(qn("w:cs"), cfg.REPORT_FONT)
    _set_rtl_run(run)
    return p


# ─── Table Helpers ────────────────────────────────────────────────────────────
def _set_table_rtl(table) -> None:
    """يضبط اتجاه الجدول RTL (متوافق مع جميع إصدارات python-docx)."""
    tbl = table._tbl
    # استخدام find/insert مباشرةً بدلاً من get_or_add_tblPr() (غير موجودة في بعض الإصدارات)
    tbl_pr = tbl.find(qn("w:tblPr"))
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)
    bidi = OxmlElement("w:bidiVisual")
    bidi.set(qn("w:val"), "1")
    tbl_pr.append(bidi)





# ─── Section Builders ────────────────────────────────────────────────────────
def _add_section_heading(doc: Document, title: str, number: int) -> None:
    """يُضيف عنوان قسم مُرقَّم بتنسيق مميّز."""
    doc.add_paragraph()   # مسافة
    p = _make_rtl_paragraph(
        doc, f"\u200F{number}. {title}",
        bold=True, size_pt=14,
        color=cfg.COLOR_PRIMARY,
    )
    # خط تحت العنوان (border)
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), cfg.COLOR_PRIMARY_HEX)
    pBdr.append(bottom)
    p._p.pPr.append(pBdr)


def _extract_section_one(ai_text: str) -> str:
    """
    يستخرج القسم الأول فقط (الملخص التنفيذي) من نص Gemini.
    يتوقف عند أول علامة تشير إلى بداية القسم الثاني.
    """
    markers = [
        "القسم الثاني", "تقييم المهام", "المهام المُنفَّذة",
        "تحليل المهام", "القسم 2",
    ]
    for marker in markers:
        idx = ai_text.find(marker)
        if idx != -1:
            return ai_text[:idx].strip()
    # لم يُعثر على فاصل — نعيد أول 40% من النص
    cutoff = int(len(ai_text) * 0.4)
    return ai_text[:cutoff].strip() or ai_text


def _add_executive_summary(doc: Document, ai_analysis: str) -> None:
    """يُضيف قسم الملخص التنفيذي من نص Gemini (القسم الأول فقط)."""
    _add_section_heading(doc, "الملخص التنفيذي", 1)

    # استخراج القسم الأول فقط — تجنّب تكرار بقية الأقسام هنا
    section_one_text = _extract_section_one(ai_analysis)

    for paragraph_text in section_one_text.split("\n"):
        stripped = paragraph_text.strip()
        # تخطّي العناوين الداخلية لتجنّب التداخل
        if stripped and not stripped.startswith("القسم الأول"):
            _make_rtl_paragraph(doc, stripped, size_pt=11, align="justify")


# Arabic ordinal words the LLM sometimes uses as task_id
_ARABIC_ORDINALS = {
    "الأول": "1", "الأولى": "1", "مهمة 1": "1", "مهمة١": "1",
    "الثاني": "2", "الثانية": "2", "مهمة 2": "2", "مهمة٢": "2",
    "الثالث": "3", "الثالثة": "3", "مهمة 3": "3", "مهمة٣": "3",
    "الرابع": "4", "الرابعة": "4", "مهمة 4": "4", "مهمة٤": "4",
    "الخامس": "5", "الخامسة": "5", "مهمة 5": "5", "مهمة٥": "5",
    "السادس": "6", "السادسة": "6", "مهمة 6": "6", "مهمة٦": "6",
    "السابع": "7", "السابعة": "7", "مهمة 7": "7", "مهمة٧": "7",
    "الثامن": "8", "الثامنة": "8", "مهمة 8": "8", "مهمة٨": "8",
    "التاسع": "9", "التاسعة": "9", "مهمة 9": "9", "مهمة٩": "9",
    "العاشر": "10", "العاشرة": "10", "مهمة 10": "10", "مهمة۱۰": "10",
    "الحادي عشر": "11", "مهمة 11": "11",
}

def _normalize(text: str) -> str:
    """Strips diacritics, spaces, and lowercases for fuzzy comparison."""
    import unicodedata
    text = text.strip().lower()
    return "".join(c for c in text
                   if unicodedata.category(c) != "Mn"  # remove tashkeel
                   and c not in (" ", "\u200f", "\u200e"))


def _extract_task_insights(ai_text: str) -> dict:
    """
    Parses the JSON block from the LLM response and builds FOUR lookup maps:
      by_id        — numeric task_id as-is  ("1", "2"...)
      by_id_arabic — Arabic ordinal task_id   (normalized to "1", "2"...)
      by_name      — original_name lowercase+strip
      by_norm      — original_name with diacritics/spaces removed
    Falls back gracefully if JSON is absent or malformed.
    """
    import re
    import json

    # Try standard fenced block first, then bare array anywhere in text
    match = (re.search(r'```json\s*(\[.*?\])\s*```', ai_text, re.DOTALL)
             or re.search(r'(\[\s*\{.*?\}\s*\])', ai_text, re.DOTALL))
    if match:
        try:
            items = json.loads(match.group(1))
            by_id        = {}
            by_id_arabic = {}
            by_name      = {}
            by_norm      = {}
            for item in items:
                raw_id   = str(item.get("task_id", "")).strip()
                raw_name = str(item.get("original_name", "")).strip()
                # numeric id ("1", "2"...)
                by_id[raw_id] = item
                # Arabic ordinal → numeric
                canonical = _ARABIC_ORDINALS.get(raw_id.lower())
                if canonical:
                    by_id_arabic[canonical] = item
                # name lookups
                by_name[raw_name.lower()] = item
                by_norm[_normalize(raw_name)] = item
            _log.debug("📊 JSON extracted: %d tasks | ids=%s | names=%s", len(items), list(by_id.keys()), list(by_name.keys()))
            return {"by_id": by_id, "by_id_arabic": by_id_arabic,
                    "by_name": by_name, "by_norm": by_norm}
        except Exception as e:
            _log.warning("⚠️  JSON parse FAILED: %s", e)
            _log.debug("📐 Raw JSON candidate: %s", match.group(1)[:300])
    else:
        _log.warning("⚠️  No JSON block found — attempting JSON repair...")
        _log.debug("📐 Response tail (last 400 chars): %s", ai_text[-400:])
        # محاولة إصلاح JSON المقطوع: ابحث عن آخر كائن مكتمل ثم أغلق المصفوفة
        last_brace = ai_text.rfind("}")
        if last_brace != -1:
            candidate = ai_text[:last_brace + 1].strip()
            bracket_open = candidate.find("[")
            if bracket_open != -1:
                candidate = candidate[bracket_open:] + "]"
            else:
                candidate = "[" + candidate + "]"
            try:
                items = json.loads(candidate)
                if isinstance(items, list) and items:
                    _log.info("🔧 JSON repaired: %d tasks recovered", len(items))
                    by_id, by_id_arabic, by_name, by_norm = {}, {}, {}, {}
                    for item in items:
                        if not isinstance(item, dict): continue
                        raw_id   = str(item.get("task_id", "")).strip()
                        raw_name = str(item.get("original_name", "")).strip()
                        by_id[raw_id] = item
                        canonical = _ARABIC_ORDINALS.get(raw_id.lower())
                        if canonical: by_id_arabic[canonical] = item
                        by_name[raw_name.lower()] = item
                        by_norm[_normalize(raw_name)] = item
                    return {"by_id": by_id, "by_id_arabic": by_id_arabic,
                            "by_name": by_name, "by_norm": by_norm}
            except Exception as repair_err:
                _log.error("❌ Repair failed: %s", repair_err)
        _log.error("❌ JSON unrecoverable.")
    return {"by_id": {}, "by_id_arabic": {}, "by_name": {}, "by_norm": {}}

def _add_tasks_section(doc: Document, office_data: OfficeData, ai_analysis: str) -> None:
    """يُضيف قسم المهام كقائمة نصية مفصلة مع خلاصة الذكاء الاصطناعي."""
    _add_section_heading(doc, "المهام التفصيلية", 2)

    tasks = office_data.get("tasks", [])
    if not tasks:
        _make_rtl_paragraph(doc, "لا توجد مهام مُسجَّلة لهذا الشهر.", size_pt=11)
        return

    insight_index = _extract_task_insights(ai_analysis)
    id_map   = insight_index["by_id"]
    name_map = insight_index["by_name"]

    for idx, task in enumerate(tasks, start=1):
        t_name  = task.get("name", "").strip()
        status  = task.get("status", "").strip()
        manager = task.get("manager", "").strip()

        # ── عنوان المهمة ──────────────────────────────────────────────────
        _make_rtl_paragraph(
            doc, f"\u200F{idx}. {t_name}",
            bold=True, size_pt=12,
            color=cfg.COLOR_PRIMARY
        )

        # ── الحالة والمسؤول ───────────────────────────────────────────────
        meta_info = []
        if status:  meta_info.append(f"الحالة: {status}")
        if manager: meta_info.append(f"المسؤول: {manager}")
        if meta_info:
            _make_rtl_paragraph(
                doc, " | ".join(meta_info),
                size_pt=10,
                color=cfg.COLOR_SECONDARY
            )

        # ── البحث عن رؤية AI — 4 طبقات متدرجة ────────────────────────
        t_name_lower = t_name.lower()
        t_name_norm  = _normalize(t_name)
        str_idx      = str(idx)

        ai_data = (
            id_map.get(str_idx)                                          # 1️⃣ numeric task_id
            or insight_index["by_id_arabic"].get(str_idx)               # 2️⃣ Arabic ordinal → numeric
            or name_map.get(t_name_lower)                                # 3️⃣ exact name (lower)
            or insight_index["by_norm"].get(t_name_norm)                 # 4️⃣ diacritic-stripped
            or next(                                                      # 5️⃣ substring fuzzy
                (v for k, v in name_map.items()
                 if t_name_lower in k or k in t_name_lower
                 or t_name_norm in _normalize(k) or _normalize(k) in t_name_norm),
                None
            )
        )

        if ai_data:
            ai_insight = ai_data.get("ai_insight", "").strip()
            _log.debug("✅ Matched task %d '‏%s' → insight %s", idx, t_name[:30], 'found' if ai_insight else '(field empty!)')
        else:
            ai_insight = ""
            _log.debug("❌ MISS task %d '‏%s' — no match in any index", idx, t_name[:30])

        # ── عرض الرؤية ────────────────────────────────────────────────────
        if ai_insight:
            p = _make_rtl_paragraph(
                doc, f"{ai_insight}",
                size_pt=10, align="justify",
                color=cfg.COLOR_TEXT
            )
            pPr = p._p.get_or_add_pPr()
            ind = OxmlElement("w:ind")
            ind.set(qn("w:right"), "360")
            pPr.append(ind)

        # ── فاصل أفقي ─────────────────────────────────────────────────────
        if idx < len(tasks):
            p = doc.add_paragraph()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "4")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "E8F0EC")
            pBdr.append(bottom)
            p._p.get_or_add_pPr().append(pBdr)
            doc.add_paragraph()


def _add_challenges_section(doc: Document, office_data: OfficeData) -> None:
    """يُضيف قسم التحديات والاختناقات مع إحصائيات المهام."""
    _add_section_heading(doc, "التحديات الإدارية", 3)

    stats = get_task_statistics(office_data)
    challenges = office_data.get("general_challenges", "")
    notes      = office_data.get("additional_notes", "")

    # إحصائيات سريعة
    stats_text = (
        f"إجمالي المهام: {stats['total']} | "
        f"المُنجَز: {stats['completed']} ({stats['completion_rate']}%) | "
        f"قيد التنفيذ: {stats['in_progress']} | "
        f"مهام لديها إشكاليات: {stats['has_issues']}"
    )
    _make_rtl_paragraph(doc, stats_text, size_pt=10,
                        color=cfg.COLOR_TEXT)

    doc.add_paragraph()
    if challenges:
        _make_rtl_paragraph(doc, "التحديات العامة:", bold=True, size_pt=11)
        _make_rtl_paragraph(doc, challenges, size_pt=11, align="justify")
    else:
        _make_rtl_paragraph(doc, "لم تُسجَّل تحديات عامة لهذا الشهر.", size_pt=11)


def _add_office_message_section(doc: Document, office_data: OfficeData) -> None:
    """يُضيف قسم رسالة المكتب إلى مجلس الإدارة (فقط إذا كان الحقل غير فارغ)."""
    notes = office_data.get("additional_notes", "").strip()
    if not notes:
        return

    _add_section_heading(doc, "رسالة المكتب إلى مجلس الإدارة", 5)
    _make_rtl_paragraph(
        doc,
        "يعرض المكتب التالي للنظر والبتت:",
        size_pt=10,
        color=cfg.COLOR_SECONDARY,
    )
    doc.add_paragraph()
    _make_rtl_paragraph(doc, notes, size_pt=11, align="justify")


def _add_audit_section(
    doc: Document,
    office_data: OfficeData,
    ai_analysis: str,
    plan_text: str,
    pdf_status: str,
) -> None:
    """
    القسم الرابع (V2.0): تحليل المطابقة بين الخطة المعتمدة والواقع التنفيذي.

    Args:
        doc:         مستند Word الجاري بناؤه.
        office_data: البيانات المُهيكَلة للمكتب.
        ai_analysis: نص التدقيق الكامل من Gemini.
        plan_text:   النص المستخرج من PDF (فارغ إذا لم يتوفر).
        pdf_status:  رسالة الحالة من pdf_handler (للعرض في التقرير).
    """
    _add_section_heading(
        doc,
        "تحليل المطابقة: الخطة المعتمدة vs الواقع التنفيذي",
        4,
    )

    has_plan = bool(plan_text and plan_text.strip())

    # ── شارة حالة PDF ──────────────────────────────────────────────────────
    if has_plan:
        badge_text  = "✅ الخطة الشهرية متوفرة — تم التدقيق المقارن"
        badge_color = cfg.COLOR_PRIMARY
    else:
        badge_text  = f"⚠️  الخطة الشهرية غير متوفرة — {pdf_status}"
        badge_color = cfg.COLOR_ACCENT

    _make_rtl_paragraph(doc, badge_text, bold=True, size_pt=11,
                        color=badge_color)
    doc.add_paragraph()

    # ── التحليل التدقيقي من Gemini ─────────────────────────────────────────
    # نبحث عن القسم الرابع في نص AI ونعرضه هنا؛ إذا لم يُميَّز نعرض النص كله
    audit_text = _extract_section_four(ai_analysis)
    for line in audit_text.split("\n"):
        stripped = line.strip()
        if stripped:
            _make_rtl_paragraph(doc, stripped, size_pt=11, align="justify")




def _extract_section_four(ai_text: str) -> str:
    """
    يحاول استخراج القسم الرابع (تحليل المطابقة) من نص Gemini.
    إذا لم يُعثر على فاصل واضح، يُعيد النص كاملًا.
    """
    # نبحث عن عناوين القسم الرابع الشائعة في مخرجات الـ AI
    markers = [
        "القسم الرابع", "تحليل المطابقة", "الخطة المعتمدة",
        "نسبة الالتزام", "الفجوات التنفيذية",
    ]
    lower = ai_text.lower()
    for marker in markers:
        idx = ai_text.find(marker)
        if idx != -1:
            return ai_text[idx:].strip()
    # لم يُعثر على فاصل — نعيد آخر 60% من النص (التوصيات عادةً في النهاية)
    cutoff = int(len(ai_text) * 0.4)
    return ai_text[cutoff:].strip() or ai_text


# ─── Cover Page ──────────────────────────────────────────────────────────────
def _add_centered_image(doc: Document, image_path: str | Path, width_inches: float = 1.8) -> None:
    """يُضيف صورة مُحاذاة للمركز مع تجاهل الخطأ إذا كان الملف غير موجود."""
    image_path = Path(image_path)
    if not image_path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))

def _add_cover(
    doc: Document,
    office_data: OfficeData,
    report_date: str,
    pdf_status: str = "",
) -> None:
    """يُضيف صفحة غلاف احترافية مع مؤشر توفر الخطة."""
    # مسافة علوية
    for _ in range(3):
        doc.add_paragraph()

    # ── لوغو المؤسسة ──────────────────────────────────────────────────────────
    logo_path = Path(__file__).parent / "assets" / "union_logo.png"
    _add_centered_image(doc, logo_path, width_inches=1.8)
    doc.add_paragraph()

    _make_rtl_paragraph(
        doc, "الاتحاد العام لطلبة سوريا",
        bold=True, size_pt=22,
        color=cfg.COLOR_PRIMARY,
        align="center",
    )
    _make_rtl_paragraph(
        doc, "منظومة المتابعة الدورية",
        bold=True, size_pt=18,
        color=cfg.COLOR_ACCENT,
        align="center",
    )
    doc.add_paragraph()
    _make_rtl_paragraph(
        doc, f"تقرير الأداء الشهري — {office_data.get('office_name', '')}",
        bold=True, size_pt=16,
        color=cfg.COLOR_TEXT,
        align="center",
    )
    _make_rtl_paragraph(
        doc, report_date,
        size_pt=13,
        color=cfg.COLOR_TEXT,
        align="center",
    )
    doc.add_paragraph()
    _make_rtl_paragraph(
        doc,
        f"مُقدَّم من: {office_data.get('submitter', '')} "
        f"| {office_data.get('submitter_phone', '')}",
        size_pt=11,
        color=cfg.COLOR_TEXT,
        align="center",
    )

    # فاصل صفحات
    doc.add_page_break()


# ─── Orchestrator-Mode Section Functions ─────────────────────────────────────
# هذه الدوال تأخذ النص الجاهز مباشرةً من ParallelOrchestrator (بدون تحليل Regex)

def _add_executive_summary_direct(doc: Document, summary_text: str) -> None:
    """يُضيف الملخص التنفيذي الجاهز مباشرةً دون استخراج."""
    _add_section_heading(doc, "الملخص التنفيذي الشامل", 1)
    if not summary_text or not summary_text.strip():
        _make_rtl_paragraph(
            doc,
            "⚠️ تعذّر توليد الملخص التنفيذي بسبب استنفاد حصة واجهات برمجة النماذج اللغوية (Gemini API) "
            "أو انقطاع الاتصال بالإنترنت. يُرجى التحقق من اتصال الشبكة وصلاحية مفتاح API الموفر.",
            size_pt=11,
        )
        return
    for para in summary_text.strip().splitlines():
        para = para.strip()
        if para:
            _make_rtl_paragraph(doc, para, size_pt=11, align="justify")


def _add_challenges_section_direct(
    doc: Document,
    office_data: OfficeData,
    challenges_text: str,
) -> None:
    """
    يُضيف قسم التحديات — يدمج الإحصائيات الخام مع النص التحليلي الجاهز من الـ AI.
    """
    _add_section_heading(doc, "التحديات الإدارية", 3)

    stats = get_task_statistics(office_data)

    # إحصائيات جدولية (كما في النسخة القديمة)
    completed   = stats.get("completed", 0)
    in_progress = stats.get("in_progress", 0)
    pending     = stats.get("pending", 0)
    total       = stats.get("total", 0)
    pct         = stats.get("completion_rate", 0.0)

    stat_lines = [
        f"📊 إجمالي المهام: {total} | مُنجزة: {completed} | "
        f"جارية: {in_progress} | معلّقة: {pending} | "
        f"نسبة الإنجاز: {pct:.0f}%"
    ]
    for line in stat_lines:
        _make_rtl_paragraph(doc, line, size_pt=10,
                            color=cfg.COLOR_PRIMARY)

    doc.add_paragraph()

    # النص التحليلي الجاهز
    if challenges_text and challenges_text.strip():
        for para in challenges_text.strip().splitlines():
            para = para.strip()
            if para:
                _make_rtl_paragraph(doc, para, size_pt=11, align="justify")
    else:
        # Fallback: عرض البيانات الخام
        raw_ch = office_data.get("general_challenges", "")
        raw_no = office_data.get("additional_notes",   "")
        if raw_ch:
            _make_rtl_paragraph(doc, f"التحديات العامة: {raw_ch}", size_pt=11)
        if raw_no:
            _make_rtl_paragraph(doc, f"الملاحظات الإضافية: {raw_no}", size_pt=11)


def _add_audit_section_direct(
    doc: Document,
    office_data: OfficeData,
    audit_text: str,
    plan_text: str,
    pdf_status: str,
) -> None:
    """يُضيف قسم تحليل المطابقة باستخدام النص الجاهز من الـ AI."""
    _add_section_heading(doc, "تحليل المطابقة: الخطة المعتمدة مقابل الإنجاز الفعلي", 4)

    # بادئة حالة الخطة
    if not (plan_text and plan_text.strip()):
        _make_rtl_paragraph(
            doc, "⚠️  لم تُرفع الخطة الشهرية — التحليل بناءً على بيانات النموذج وحدها.",
            size_pt=10, color=cfg.COLOR_ACCENT
        )

    doc.add_paragraph()

    if audit_text and audit_text.strip():
        for para in audit_text.strip().splitlines():
            para = para.strip()
            if para:
                _make_rtl_paragraph(doc, para, size_pt=11, align="justify")
    else:
        _make_rtl_paragraph(doc, "لم يتوفر تحليل المطابقة.", size_pt=11)


def _convert_to_pdf(docx_path: Path) -> Path | None:
    """
    يحوّل ملف .docx إلى .pdf بطريقة متوافقة مع أنظمة Windows و Linux.
    يتحقق من وجود محرك التحويل (LibreOffice) قبل التنفيذ لضمان استقرار النظام.
    """
    system = platform.system()
    binary = None

    # ── 1. تحديد مسار محرك التحويل ───────────────────────────────────────────
    if system == "Linux":
        binary = shutil.which("libreoffice")
    elif system == "Windows":
        # البحث في PATH أولاً
        binary = shutil.which("soffice")
        if not binary:
            # مسارات التثبيت الشائعة لـ LibreOffice على Windows
            common_paths = [
                r"C:\Program Files\LibreOffice\program\soffice.exe",
                r"C:\Program Files (x86)\LibreOffice\program\soffice.exe"
            ]
            for p in common_paths:
                if Path(p).exists():
                    binary = p
                    break

    # ── 2. التحقق من توفر المحرك ─────────────────────────────────────────────
    if not binary:
        _log.warning("⚠️ PDF Conversion Engine (LibreOffice/soffice) not found. "
                     "Skipping PDF generation, Word document is preserved.")
        return None

    try:
        pdf_path = docx_path.with_suffix(".pdf")
        _log.info("🔄 Converting to PDF (%s): %s...", system, pdf_path.name)
        
        # ── 3. تنفيذ أمر التحويل ────────────────────────────────────────────────
        cmd = [
            binary,
            "--headless",
            "--convert-to", "pdf",
            "--outdir", str(docx_path.parent),
            str(docx_path)
        ]
        
        subprocess.run(
            cmd,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            timeout=40,  # مهلة زمنية كافية للتحويل
            check=True
        )
        
        if pdf_path.exists():
            _log.info("✅ PDF Generated: %s", pdf_path.name)
            return pdf_path
            
    except Exception as e:
        _log.warning("⚠️ PDF Conversion failed (non-critical): %s", e)
    
    return None


# ─── Main Builder ─────────────────────────────────────────────────────────────
def build_report(
    office_data: OfficeData,
    ai_analysis,                   # str (legacy) | dict (orchestrator)
    output_path: str | Path,
    plan_text: str = "",
    pdf_status: str = "",
) -> Path:
    """
    يُنشئ ملف Word (.docx) احترافيًا باللغة العربية.

    Args:
        office_data:  البيانات المُهيكَلة من data_parser.
        ai_analysis:  نص التدقيق (str) أو قاموس النتائج الموازية (dict).
        output_path:  مسار ملف الـ .docx المُراد إنشاؤه.
        plan_text:    النص المستخرج من PDF الخطة (اختياري).
        pdf_status:   رسالة حالة PDF من pdf_handler (اختياري).

    Returns:
        مسار الملف المُنشَأ كـ Path object.
    """
    output_path = Path(output_path)

    # ── تفريق المسارين: orchestrator dict / legacy string ────────────────────
    if isinstance(ai_analysis, dict):
        # وضع الـ Orchestrator: كل مفتاح يحمل نص قسمه مباشرةً
        _summary_text    = ai_analysis.get("summary",    "") or ""
        _tasks_text      = ai_analysis.get("tasks",      "") or ""
        _audit_text      = ai_analysis.get("audit",      "") or ""
        _challenges_text = ai_analysis.get("challenges", "") or ""
        _legacy_str      = None    # لا يُستخدم
    else:
        # وضع Legacy: استخراج الأقسام من النص الكامل
        _legacy_str      = ai_analysis
        _summary_text    = None
        _tasks_text      = None
        _audit_text      = None
        _challenges_text = None

    doc = Document()

    # ── إعدادات الصفحة ──
    section = doc.sections[0]
    section.page_width  = Inches(8.27)   # A4 عرض
    section.page_height = Inches(11.69)  # A4 ارتفاع
    section.left_margin  = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin   = Inches(0.8)
    section.bottom_margin= Inches(0.8)

    # تاريخ التقرير مسحب من مدخلات المستخدم (target_month_name) أو الـ Fallback
    current_year = datetime.now().year
    month_name = office_data.get("target_month_name", "")
    if month_name:
        report_date = f"شهر {month_name} {current_year}"
    else:
        report_date = datetime.now().strftime("%B %Y")

    # ── صفحة الغلاف ──
    _add_cover(doc, office_data, report_date, pdf_status)

    # ── الأقسام — orchestrator: نص جاهز / legacy: استخراج من النص الكامل ──
    if _legacy_str is not None:
        _add_executive_summary(doc, _legacy_str)
        _add_tasks_section(doc, office_data, _legacy_str)
        _add_challenges_section(doc, office_data)
        _add_audit_section(doc, office_data, _legacy_str, plan_text, pdf_status)
    else:
        _add_executive_summary_direct(doc, _summary_text)
        _add_tasks_section(doc, office_data, _tasks_text)
        _add_challenges_section_direct(doc, office_data, _challenges_text)
        _add_audit_section_direct(doc, office_data, _audit_text, plan_text, pdf_status)

    _add_office_message_section(doc, office_data)

    # ── التذييل ──
    doc.add_paragraph()
    _make_rtl_paragraph(
        doc,
        "— وثيقة مُولَّدة بواسطة منظومة المتابعة الدورية · الاتحاد العام لطلبة سوريا —",
        size_pt=9,
        color=RGBColor(0xDD, 0xB5, 0x57),   # ذهبي المؤسسة
        align="center",
    )

    doc.save(output_path)
    _log.info("📄 Report saved: %s", output_path.name)

    # ── V2.2: تصدير نسخة PDF مرآتية ──────────────────────────────────────────
    _convert_to_pdf(output_path)

    return output_path


# ─── Quick Test ──────────────────────────────────────────────────────────────
if __name__ == "__main__":
    data = {
        "office_name": "مكتب دمشق",
        "submitter": "أحمد علي",
        "submitter_phone": "0900000001",
        "target_month_name": "كانون الثاني",
        "target_month_num": 1,
        "monthly_plan_link": "http://plan.pdf",
        "general_challenges": "نقص في الميزانية",
        "additional_notes": "نأمل في زيادة الدعم",
        "tasks": [
            {
                "manager": "محمد سالم",
                "manager_phone": "0911111111",
                "name": "تنظيم ندوة",
                "description": "ندوة حول الحقوق الطلابية",
                "type": "ثقافي",
                "mechanism": "حضوري",
                "status": "مكتمل",
                "issues": "",
                "file_link": ""
            }
        ]
    }
    cfg.REPORTS_DIR.mkdir(exist_ok=True)
    path = build_report(
        data,
        "هذا ملخص تدقيقي تجريبي من Gemini V2.0.",
        cfg.REPORTS_DIR / "test_report_v2.docx",
        plan_text="الخطة الشهرية: 1. ندوة 2. ورشة عمل 3. اجتماع",
        pdf_status="",
    )
    _log.info("✅ Test report V2.0: %s", path)
